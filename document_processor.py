"""
Извлечение текста из документов (PDF, DOC, DOCX) — ЛОКАЛЬНО через PyMuPDF
"""

import re
import os
from typing import List, Dict

# MIME-типы
MIME_TYPES = {
    ".pdf": "application/pdf",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
}

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx"}


def get_file_type(file_path: str) -> str:
    """Возвращает тип файла."""
    ext = os.path.splitext(file_path)[1].lower()
    return ext.lstrip(".")


def is_supported(file_path: str) -> bool:
    """Проверяет, поддерживается ли формат."""
    ext = os.path.splitext(file_path)[1].lower()
    return ext in SUPPORTED_EXTENSIONS


# =====================================================
# ИЗВЛЕЧЕНИЕ ТЕКСТА ИЗ PDF (PyMuPDF)
# =====================================================
def extract_text_from_pdf(file_path: str) -> str:
    """Извлекает текст из PDF через PyMuPDF."""
    try:
        import fitz
    except ImportError:
        raise RuntimeError("PyMuPDF не установлен: pip install PyMuPDF")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    doc = fitz.open(file_path)
    pages_text = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            pages_text.append(text)

    doc.close()
    
    print(f"   ✓ PDF: {len(pages_text)} страниц")
    return "\n\n".join(pages_text)


# =====================================================
# ИЗВЛЕЧЕНИЕ ТЕКСТА ИЗ DOCX
# =====================================================
def extract_text_from_docx(file_path: str) -> str:
    """Извлекает текст из DOCX."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx не установлен: pip install python-docx")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    doc = Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                paragraphs.append(" | ".join(row_text))

    print(f"   ✓ DOCX: {len(paragraphs)} параграфов")
    return "\n\n".join(paragraphs)


# =====================================================
# ИЗВЛЕЧЕНИЕ ТЕКСТА ИЗ DOC (старый формат)
# =====================================================
def extract_text_from_doc(file_path: str) -> str:
    """Извлекает текст из DOC (пробует несколько способов)."""
    import subprocess

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    # Способ 1: antiword
    try:
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0 and result.stdout.strip():
            print(f"   ✓ DOC: извлечено через antiword")
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Способ 2: catdoc
    try:
        result = subprocess.run(
            ["catdoc", "-w", file_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0 and result.stdout.strip():
            print(f"   ✓ DOC: извлечено через catdoc")
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Способ 3: попробовать как DOCX
    try:
        return extract_text_from_docx(file_path)
    except:
        pass

    raise RuntimeError(
        "Не удалось извлечь текст из DOC. Установите antiword:\n"
        "  apt-get install antiword"
    )


# =====================================================
# УНИВЕРСАЛЬНАЯ ФУНКЦИЯ
# =====================================================
def extract_text(file_path: str, use_api: bool = False) -> str:
    """
    Извлекает текст из документа (PDF, DOC, DOCX).
    use_api игнорируется — всегда локально.
    """
    file_type = get_file_type(file_path)

    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    elif file_type == "doc":
        return extract_text_from_doc(file_path)
    else:
        raise ValueError(f"Неподдерживаемый формат: {file_path}")


# =====================================================
# ОЧИСТКА И ЧАНКИНГ
# =====================================================
def clean_text(text: str) -> str:
    """Очистка текста."""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\n +", "\n", text)
    text = re.sub(r"[\x00-\x09\x0b-\x0c\x0e-\x1f]", "", text)
    return text.strip()


def split_into_chunks(
    text: str,
    chunk_size: int = 800,
    chunk_overlap: int = 200,
) -> List[Dict]:
    """Разбивает текст на чанки с перекрытием."""
    text = clean_text(text)

    if not text:
        return []

    if len(text) <= chunk_size:
        return [{"text": text, "chunk_id": 0}]

    chunks = []
    start = 0
    chunk_id = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))

        if end < len(text):
            search_from = start + int(chunk_size * 0.7)
            best_break = -1

            for sep in [".\n", ". ", "!\n", "! ", "?\n", "? ", ";\n", "; ", "\n\n", "\n"]:
                pos = text.rfind(sep, search_from, end)
                if pos != -1 and pos + len(sep) > best_break:
                    best_break = pos + len(sep)

            if best_break > search_from:
                end = best_break

        chunk_text = text[start:end].strip()

        if chunk_text and len(chunk_text) >= 50:
            chunks.append({"text": chunk_text, "chunk_id": chunk_id})
            chunk_id += 1

        next_start = end - chunk_overlap
        if next_start <= start:
            next_start = end
        start = next_start

    return chunks


# =====================================================
# ГЛАВНАЯ ФУНКЦИЯ
# =====================================================
def process_document(
    file_path: str,
    chunk_size: int = 800,
    chunk_overlap: int = 200,
    use_api: bool = False,  # игнорируется
) -> List[Dict]:
    """
    Полный пайплайн: документ → текст → чанки.
    Поддерживает: PDF, DOC, DOCX
    """
    filename = os.path.basename(file_path)
    file_type = get_file_type(file_path)

    print(f"📄 Обработка: {filename} [{file_type.upper()}]")

    if not is_supported(file_path):
        print(f"   ❌ Неподдерживаемый формат")
        return []

    try:
        text = extract_text(file_path)
    except Exception as e:
        print(f"   ❌ Ошибка: {e}")
        return []

    print(f"   Извлечено символов: {len(text)}")

    if not text.strip():
        print("   ⚠️ Документ не содержит текста")
        return []

    chunks = split_into_chunks(text, chunk_size, chunk_overlap)
    print(f"   Создано чанков: {len(chunks)}")

    if chunks:
        preview = chunks[0]["text"][:120].replace("\n", " ")
        print(f"   Превью: «{preview}...»")

    return chunks


# Алиас
def process_pdf(file_path: str, chunk_size: int = 800, chunk_overlap: int = 200, use_api: bool = False) -> List[Dict]:
    return process_document(file_path, chunk_size, chunk_overlap, use_api)